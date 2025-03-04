# 가장 기본적인 기능(문서 열기, 저장, 글자 쓰기 등등)
import docx
from docx import Document
from docx.shared import Pt
from docx.shared import Cm, Inches
from docx.shared import RGBColor
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH

from datetime import datetime
import configparser
import subprocess
import os


# 단순 띄어쓰기
def add_new_line(doc, paragraph_num):
    paragraph = doc.paragraphs[paragraph_num]
    paragraph.add_run('\n')
    
# 여러줄 띄어쓰기
def iter_new_line(doc, paragraph_num, iter):
    for _ in range(iter):
        add_new_line(doc, paragraph_num)

# 새로운 섹션 추가
def add_new_section(doc, section_num, left=1.0, right=1.0, top=1.0, bottom=1.0):
    # 새로운 색션에서 처음부터 시작
    doc.add_section(start_type=WD_SECTION.NEW_PAGE)
    
    # 섹션 넘버 지정 필요
    section = doc.sections[section_num]

    # 여백 지정
    section.left_margin = Inches(left)  # 왼쪽 여백
    section.right_margin = Inches(right)  # 오른쪽 여백
    section.top_margin = Inches(top)  # 위쪽 여백
    section.bottom_margin = Inches(bottom)  # 아래쪽 여백
    

def add_text_and_image(doc, text, image_path):
    # Add a table with one row and two cells
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False  # Disable automatic cell width adjustment


    # Add the image to the first cell
    cell_0_0 = table.cell(0, 0)
    cell_0_0.paragraphs[0].add_run(text)
    # Add the text to the second cell
    cell_0_1 = table.cell(0, 1)
    cell_0_1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cell_0_1.paragraphs[0].add_run().add_picture(image_path, width=Inches(1.4))


def attach_image(doc, solution_num, image_folder_path):


    if solution_num == 0: # 12-1번
        image_path = os.path.join(image_folder_path, "11.png")
        para = doc.add_paragraph() # 이미지 문단
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 가운데 정렬
        para.add_run().add_picture(image_path, width=Cm(15))
    elif solution_num == 1: # 12-2번
        image_path = os.path.join(image_folder_path, "1.png")
        para = doc.add_paragraph() # 이미지 문단
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 가운데 정렬
        para.add_run().add_picture(image_path, width=Cm(15))


    elif solution_num == 2: # 13-1번
        image_path = os.path.join(image_folder_path, "1.png")
        para = doc.add_paragraph() # 이미지 문단
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 가운데 정렬
        para.add_run().add_picture(image_path, width=Cm(15))
    elif solution_num == 3: # 13-2번
        image_path = os.path.join(image_folder_path, "2.png")
        para = doc.add_paragraph() # 이미지 문단
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 가운데 정렬
        para.add_run().add_picture(image_path, width=Cm(15))
    elif solution_num == 4: # 13-3번
        image_path = os.path.join(image_folder_path, "1.png")
        para = doc.add_paragraph() # 이미지 문단
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 가운데 정렬
        para.add_run().add_picture(image_path, width=Cm(15))

        image_path = os.path.join(image_folder_path, "3.png")
        para = doc.add_paragraph() # 이미지 문단
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 가운데 정렬
        para.add_run().add_picture(image_path, width=Cm(15))

        image_path = os.path.join(image_folder_path, "6.png")
        para = doc.add_paragraph() # 이미지 문단
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 가운데 정렬
        para.add_run().add_picture(image_path, width=Cm(15))
    elif solution_num == 5: # 13-4번
        image_path = os.path.join(image_folder_path, "5.png")
        para = doc.add_paragraph() # 이미지 문단
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 가운데 정렬
        para.add_run().add_picture(image_path, width=Cm(15))
    elif solution_num == 6: # 13-5번
        image_path = os.path.join(image_folder_path, "1.png")
        para = doc.add_paragraph() # 이미지 문단
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 가운데 정렬
        para.add_run().add_picture(image_path, width=Cm(15))
    elif solution_num == 7: # 13-6번
        image_path = os.path.join(image_folder_path, "3.png")
        para = doc.add_paragraph() # 이미지 문단
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 가운데 정렬
        para.add_run().add_picture(image_path, width=Cm(15))

        image_path = os.path.join(image_folder_path, "4.png")
        para = doc.add_paragraph() # 이미지 문단
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 가운데 정렬
        para.add_run().add_picture(image_path, width=Cm(15))


    elif solution_num == 8: # 14-1번
        image_path = os.path.join(image_folder_path, "1.png")
        para = doc.add_paragraph() # 이미지 문단
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 가운데 정렬
        para.add_run().add_picture(image_path, width=Cm(15))

        image_path = os.path.join(image_folder_path, "8.png")
        para = doc.add_paragraph() # 이미지 문단
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 가운데 정렬
        para.add_run().add_picture(image_path, width=Cm(15))
    elif solution_num == 9: # 14-2번
        image_path = os.path.join(image_folder_path, "10.png")
        para = doc.add_paragraph() # 이미지 문단
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 가운데 정렬
        para.add_run().add_picture(image_path, width=Cm(15))

        image_path = os.path.join(image_folder_path, "9.png")
        para = doc.add_paragraph() # 이미지 문단
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 가운데 정렬
        para.add_run().add_picture(image_path, width=Cm(15))
    elif solution_num == 10: # 14-3번
        image_path = os.path.join(image_folder_path, "2.png")
        para = doc.add_paragraph() # 이미지 문단
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 가운데 정렬
        para.add_run().add_picture(image_path, width=Cm(15))

        image_path = os.path.join(image_folder_path, "11.png")
        para = doc.add_paragraph() # 이미지 문단
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 가운데 정렬
        para.add_run().add_picture(image_path, width=Cm(15))
    elif solution_num == 12: # 14-5번 (14-4 생략)
        para = doc.add_paragraph() # 이미지 문단
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 가운데 정렬
        image_path = os.path.join(image_folder_path, "8.png")
        para.add_run().add_picture(image_path, width=Cm(15))
    else:
        pass


    

def make_pdf(user_name, date, recommend_indice, txt_save_path = '.', docx_save_path = './report.docx', pdf_save_path = '.'):

    # 설정파일 읽기
    config = configparser.ConfigParser()    
    config.read('./config.ini', encoding='utf-8')
    image_folder_path = config['makepdf']['image_path']


    # 새로운 문서 생성
    doc = Document()

    ############################## 표지 ##############################
    name = user_name
    main_title = name + "님\n낙상예방 9988\n보고서"

    # 첫페이지 상단 공백
    first_section = doc.sections[0]
    first_section.left_margin = Inches(0.7)
    first_section.bottom_margin = Pt(0)

    doc.add_paragraph()
    iter_new_line(doc, 0, 2)

    doc.add_paragraph(main_title)
    para = doc.paragraphs[1].runs

    for run in para:
        run.font.size = Pt(40)
        run.bold = True

    # 조사 날짜 get
    current_time = date
    current_year = current_time.year
    current_month = current_time.month
    current_day = current_time.day

    reporting_time_string = current_time.strftime("%Y.%m.%d") 
    main_second = " Fall Prevention Report: " + reporting_time_string
    doc.add_paragraph(main_second)
    para = doc.paragraphs[2].runs

    for run in para:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor.from_string('9C9393')

    main_third = " " + name + f"님의 신체측정과 주거측정을 통해 낙상위험정도를 파악하기 위해\n  {current_year}년 {current_month}월 {current_day}일에 작성된 보고서 입니다."
    doc.add_paragraph(main_third)
    para = doc.paragraphs[3].runs

    for run in para:
        run.font.size = Pt(10)

    main_forth = "  낙상예방 9988은 정확한 측정을 통해 낙상예방 솔루션을 제시해 드립니다."
    doc.add_paragraph(main_forth)
    para = doc.paragraphs[4].runs

    for run in para:
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor.from_string('9C9393')

    # 메인 페이지 이미지 받아오기    
    image_path = os.path.join(image_folder_path, "main_page_image.JPG")
    para = doc.add_paragraph() # 이미지 문단
    iter_new_line(doc, 5, 3)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 가운데 정렬
    para.add_run().add_picture(image_path, width=Cm(17.5))  # Adjust the width as needed
        
        
    # 새로운 섹션 시작
    add_new_section(doc, 1)
    ############################## 2페이지부터 진단 결과 ##############################
    try: # docx를 pdf로 변환
        file_open_path = txt_save_path
        with open(file_open_path, "r", encoding='utf-8') as file:
            content = file.read()
        
        # 텍스트 파일 내용을 docx에 추가합니다.
        lines = content.split('\n')
        for line in lines:
            if line.startswith("[") and line.endswith("]"):
                title = line[1:-1]

                if title == '기본 신체 정보':
                    image_path = os.path.join(image_folder_path, "body_title.jpg")
                    para = doc.add_paragraph() # 이미지 문단
                    para.add_run().add_picture(image_path, width=Cm(16))
                if title == '기본 주거 정보':
                    doc.add_page_break()
                    image_path = os.path.join(image_folder_path, "house_title.jpg")
                    para = doc.add_paragraph() # 이미지 문단
                    para.add_run().add_picture(image_path, width=Cm(16))
                if title == '종합 평가':
                    doc.add_page_break()
                    image_path = os.path.join(image_folder_path, "summary_title.jpg")
                    para = doc.add_paragraph() # 이미지 문단
                    para.add_run().add_picture(image_path, width=Cm(16))
                    doc.add_heading("생활환경평가", level=1)
                if title == '신체 정보':
                    image_path = os.path.join(image_folder_path, "disease_image.JPG")
                    para = doc.add_paragraph() # 이미지 문단
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    para.add_run().add_picture(image_path, width=Cm(12.5))
                if title == '운동 주기':
                    image_path = os.path.join(image_folder_path, "body_image.png")
                    para = doc.add_paragraph() # 이미지 문단
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    para.add_run().add_picture(image_path, width=Cm(12))


                if title != '종합 평가':
                    doc.add_heading(title, level=1) # 제목 넣기

            elif line.startswith("<") and line.endswith(">"):
                doc.add_paragraph("\n")
                doc.add_heading(line, level=2)
            elif line == "": # 아무 내용이 없을때...
                pass
            else:
                para = doc.add_paragraph(line)

    except FileNotFoundError:
        print('파일을 찾을 수 없습니다.')


    # 최종 솔루션 이미지 보이기
    attach_image(doc, recommend_indice[0], image_folder_path)


    ############################## 마지막 페이지부터 진단 결과 ##############################

    add_new_section(doc, 2) # 새로운 섹션 추가
    last_section = doc.sections[2] # 아래 공백 0
    last_section.bottom_margin = Pt(0)

    image_path = os.path.join(image_folder_path, "contact_title.jpg")
    para = doc.add_paragraph("\n\n\n") # 줄 띄우기용
    para = doc.add_paragraph() # 이미지 문단
    para.add_run().add_picture(image_path, width=Cm(16))
    doc.add_paragraph("\n본 보고서는 인공지능(ChatGPT)을 활용하여 만든 낙상예방 솔루션으로서 해피에이징에 저작권이 있습니다. 보고서에 관하여 궁금하거나 문의하실 내용이 있으시면 아래 메일로 연락 주시기 바랍니다.\n\n메일주소: happy@happy-aging.co.kr\n홈페이지: https://happy-aging.co.kr")


    # 바닥글 삽입
    section = doc.sections[0]
    footer = section.footer
    footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    footer_paragraph = footer.paragraphs[0].add_run('해피에이징 낙상 예방 9988')
    footer_paragraph.font.size = Pt(8)
    footer_paragraph.font.color.rgb = RGBColor(128, 128, 128)


    doc.save(docx_save_path)

    subprocess.call(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', pdf_save_path, docx_save_path])